"""
Text + table extraction from syllabus files: PDF, DOCX, images, plain text.

Syllabi are very often multi-column schedule TABLES. Flattening them to plain
text interleaves the columns into garbage, so we extract tables structurally
when we can and let the parser walk rows/cells. Plain text is kept as a
fallback and for the LLM path.

Heavy libraries are imported lazily so this module loads even when they aren't
installed (e.g. for unit-testing the date logic).
"""
from __future__ import annotations

import io
import os


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


# --- PDF ---------------------------------------------------------------

def _pdf(data: bytes) -> tuple[str, list[list[list[str]]]]:
    import pdfplumber  # lazy

    text_parts: list[str] = []
    tables: list[list[list[str]]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
            for t in page.extract_tables() or []:
                # normalize None cells to ""
                tables.append([[(c or "") for c in row] for row in t])

    text = "\n".join(text_parts).strip()
    if len(text) < 40 and not tables:
        try:
            text = _ocr_pdf(data)
        except Exception:
            pass
    return text, tables


def _ocr_pdf(data: bytes) -> str:
    from pdf2image import convert_from_bytes  # lazy
    import pytesseract  # lazy

    images = convert_from_bytes(data)
    return "\n".join(pytesseract.image_to_string(img) for img in images)


# --- DOCX --------------------------------------------------------------

def _docx(data: bytes) -> tuple[str, list[list[list[str]]]]:
    import docx  # python-docx, lazy

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
            parts.append(" | ".join(c for c in cells if c))
        tables.append(rows)
    return "\n".join(parts).strip(), tables


# --- images ------------------------------------------------------------

def _image(data: bytes) -> str:
    from PIL import Image  # lazy
    import pytesseract  # lazy

    return pytesseract.image_to_string(Image.open(io.BytesIO(data))).strip()


# --- public API --------------------------------------------------------

def extract(filename: str, data: bytes) -> dict:
    """Return {'text': str, 'tables': list[table]} for any supported file."""
    ext = _ext(filename)
    if ext == ".pdf":
        text, tables = _pdf(data)
        return {"text": text, "tables": tables}
    if ext in (".docx", ".doc"):
        text, tables = _docx(data)
        return {"text": text, "tables": tables}
    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"):
        return {"text": _image(data), "tables": []}
    return {"text": data.decode("utf-8", errors="ignore"), "tables": []}


def extract_text(filename: str, data: bytes) -> str:
    """Backwards-compatible plain-text extraction."""
    return extract(filename, data)["text"]
